import os
import csv
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_UNDERLINE
from docx.enum.section import WD_SECTION, WD_ORIENTATION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn, nsdecls

def create_element(name):
    """Create an XML element."""
    return OxmlElement(name)

def create_attribute(element, name, value):
    """Add an attribute to an XML element."""
    element.set(qn(name), value)

def add_hyperlink(paragraph, text, anchor, color='0563C1'):
    """Add a hyperlink to a paragraph with proper formatting."""
    # Add hyperlink
    hyperlink_element = parse_xml(r'<w:hyperlink w:anchor="{anchor}" xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>'.format(anchor=anchor))
    
    # Create run
    run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    
    # Style the hyperlink
    rStyle = OxmlElement('w:rStyle')
    rStyle.set(qn('w:val'), 'Hyperlink')
    rPr.append(rStyle)
    
    # Add color
    c = OxmlElement('w:color')
    c.set(qn('w:val'), color)
    rPr.append(c)
    
    # Add underline
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)
    
    run.append(rPr)
    
    # Add text
    t = OxmlElement('w:t')
    t.text = text
    run.append(t)
    
    hyperlink_element.append(run)
    paragraph._element.append(hyperlink_element)

def add_bookmark(paragraph, bookmark_name):
    """Add a bookmark to a paragraph."""
    bookmark_id = str(abs(hash(bookmark_name)) % 1000000)
    
    bookmark_start = OxmlElement('w:bookmarkStart')
    bookmark_start.set(qn('w:id'), bookmark_id)
    bookmark_start.set(qn('w:name'), bookmark_name)
    
    bookmark_end = OxmlElement('w:bookmarkEnd')
    bookmark_end.set(qn('w:id'), bookmark_id)
    
    paragraph._p.insert(0, bookmark_start)
    paragraph._p.append(bookmark_end)

def set_cell_background(cell, color):
    """Set background color of a table cell."""
    cell_properties = cell._element.get_or_add_tcPr()
    shading = OxmlElement('w:shd')
    shading.set(qn('w:val'), 'clear')
    shading.set(qn('w:color'), 'auto')
    shading.set(qn('w:fill'), color)
    cell_properties.append(shading)

def set_cell_borders(cell, top=None, bottom=None, left=None, right=None):
    """Set custom border colors for a table cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    
    # Remove existing borders
    tcBorders = tcPr.find(qn('w:tcBorders'))
    if tcBorders is not None:
        tcPr.remove(tcBorders)
    
    # Create new borders element
    tcBorders = OxmlElement('w:tcBorders')
    
    # Helper function to create border
    def create_border(border_type, color, size='4'):
        border = OxmlElement(f'w:{border_type}')
        if color == 'none':
            border.set(qn('w:val'), 'nil')
        else:
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), size)
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), color)
        return border
    
    # Add borders as specified - default to none if not specified
    tcBorders.append(create_border('top', top if top else 'none'))
    tcBorders.append(create_border('bottom', bottom if bottom else 'none'))
    tcBorders.append(create_border('left', left if left else 'none'))
    tcBorders.append(create_border('right', right if right else 'none'))
    
    tcPr.append(tcBorders)

def set_table_borders(table):
    """Set table-level borders to ensure consistent grid appearance."""
    tbl = table._tbl
    
    # Get or create table properties
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    
    # Remove any existing table borders
    tblBorders = tblPr.find(qn('w:tblBorders'))
    if tblBorders is not None:
        tblPr.remove(tblBorders)
    
    # Create new table borders element
    tblBorders = OxmlElement('w:tblBorders')
    
    # Define border properties
    border_attrs = {
        'w:val': 'single',
        'w:sz': '4',
        'w:space': '0',
        'w:color': 'CCCCCC'
    }
    
    # Add borders for all sides
    for border_name in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{border_name}')
        for attr_name, attr_value in border_attrs.items():
            border.set(qn(attr_name), attr_value)
        tblBorders.append(border)
    
    tblPr.append(tblBorders)

def parse_allowance_csv(csv_path):
    """Parse the allowance CSV file and return structured data."""
    allowances = {}
    classes = set()
    
    with open(csv_path, 'r', encoding='utf-8') as file:
        reader = csv.DictReader(file)
        
        for row in reader:
            class_name = row['class'].strip()
            age_from = row['ageFrom'].strip()
            
            # Add class to set
            classes.add(class_name)
            
            # Initialize class in allowances if not exists
            if class_name not in allowances:
                allowances[class_name] = {}
            
            # Store allowance data for this age
            age_key = age_from if age_from != '64' else '64+'
            
            allowances[class_name][age_key] = {
                'EE': row.get('EE', '').strip(),
                'ES': row.get('ES', '').strip(),
                'EC1': row.get('EC1', '').strip(),
                'EC2': row.get('EC2', '').strip(),
                'ECmax': row.get('ECmax', '').strip(),
                'FA1': row.get('FA1', '').strip(),
                'FA2': row.get('FA2', '').strip(),
                'FAmax': row.get('FAmax', '').strip()
            }
    
    return sorted(list(classes)), allowances

def format_currency(value):
    """Format currency value without decimals."""
    if not value or value == '':
        return ''
    
    try:
        # Remove any existing formatting
        clean_value = value.replace('$', '').replace(',', '').strip()
        
        if not clean_value:
            return ''
            
        # Convert to float then int to remove decimals
        amount = int(float(clean_value))
        
        # Format with comma separators and dollar sign
        return f'${amount:,}'
    except:
        return value

def create_ichra_document(output_filename='ICHRA_Allowance_Model_2025.docx', 
                         header_image_path=None, 
                         csv_path=None,
                         states=None):
    """Create the ICHRA Monthly Employer Contributions document.
    
    Args:
        output_filename: Name of the output DOCX file
        header_image_path: Path to the header image file
        csv_path: Path to the allowance CSV file
        states: List of state codes (will be auto-detected from CSV if csv_path is provided)
    """
    
    # If CSV path is provided, parse it to get classes and data
    allowance_data = {}
    if csv_path and os.path.exists(csv_path):
        states, allowance_data = parse_allowance_csv(csv_path)
        print(f"Found {len(states)} classes in CSV: {', '.join(states)}")
    elif states is None:
        states = []
    
    # Create a new Document
    doc = Document()
    
    # Set page size to Letter in landscape orientation (11" x 8.5")
    section = doc.sections[0]
    section.page_height = Inches(8.5)
    section.page_width = Inches(11)
    section.orientation = WD_ORIENTATION.LANDSCAPE
    
    # Set margins
    section.top_margin = Inches(0)  # Remove top margin completely
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)
    section.header_distance = Inches(0)  # No distance between header and top of page
    
    # Configure header - just add the image directly
    if header_image_path and os.path.exists(header_image_path):
        header = section.header
        header.is_linked_to_previous = False
        
        # Clear any default content
        for paragraph in header.paragraphs:
            p = paragraph._element
            p.getparent().remove(p)
            
        # Add a new paragraph
        header_para = header.add_paragraph()
        
        # Access paragraph properties to set negative indents
        pPr = header_para._p.get_or_add_pPr()
        
        # Create indent element with negative values to extend beyond margins
        ind = OxmlElement('w:ind')
        ind.set(qn('w:left'), str(int(-0.75 * 1440)))  # -0.75 inches in twips
        ind.set(qn('w:right'), str(int(-0.75 * 1440)))  # -0.75 inches in twips
        ind.set(qn('w:firstLine'), '0')
        pPr.append(ind)
        
        # Remove any spacing
        spacing = OxmlElement('w:spacing')
        spacing.set(qn('w:before'), '0')
        spacing.set(qn('w:after'), '0')
        spacing.set(qn('w:line'), '240')  # Single line spacing
        spacing.set(qn('w:lineRule'), 'auto')
        pPr.append(spacing)
        
        # Set alignment
        header_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        # Add the image with exact width to cover full page including margins
        run = header_para.add_run()
        run.add_picture(header_image_path, width=Inches(11.5))  # Slightly wider to ensure full coverage
        
        # Adjust first section to add back top margin for content
        section.top_margin = Inches(0.5)
    else:
        if header_image_path:
            print(f"Warning: Header image not found at path: {header_image_path}")
        else:
            print("Warning: No header image path provided")
    
    # Main document content with tighter spacing
    
    # Add title
    title_para = doc.add_paragraph()
    title_para.style = 'Normal'  # Use Normal style for better control
    title_run = title_para.add_run('ICHRA Monthly Employer Contributions')
    title_run.font.size = Pt(16)
    title_run.font.bold = True
    title_run.font.name = 'Calibri'
    title_para.space_before = Pt(6)
    title_para.space_after = Pt(6)
    
    # Add Class heading
    class_para = doc.add_paragraph()
    class_para.style = 'Normal'
    class_run = class_para.add_run('Class')
    class_run.font.size = Pt(14)
    class_run.font.bold = True
    class_run.font.name = 'Calibri'
    class_para.space_before = Pt(6)
    class_para.space_after = Pt(6)
    
    # Add hyperlinks only if states are provided
    if states:
        for i, state in enumerate(states):
            state_para = doc.add_paragraph()
            # Create anchor name - replace spaces and special chars with underscores
            anchor_name = state.lower().replace(' ', '_').replace('-', '_')
            add_hyperlink(state_para, state, anchor_name)
            
            # Add spacing after each link except the last one
            if i < len(states) - 1:
                state_para.space_after = Pt(6)
            else:
                state_para.space_after = Pt(12)
    
    # Function to create a state table
    def create_state_table(state_name, anchor_name, add_page_break=False):
        # Add page break before state if requested
        if add_page_break:
            doc.add_page_break()
        
        # Add state heading with bookmark (no spacing paragraph before it)
        state_heading = doc.add_paragraph()
        state_heading.style = 'Normal'
        add_bookmark(state_heading, anchor_name)
        state_run = state_heading.add_run(state_name)
        state_run.font.size = Pt(13)
        state_run.font.bold = True
        state_run.font.name = 'Calibri'
        state_heading.space_before = Pt(6)
        state_heading.space_after = Pt(6)
        
        # Create table
        table = doc.add_table(rows=48, cols=9)
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        table.autofit = False
        
        # Enable header row repetition on each page
        header_row = table.rows[0]
        tr = header_row._tr
        trPr = tr.get_or_add_trPr()
        tblHeader = OxmlElement('w:tblHeader')
        trPr.append(tblHeader)
        
        # Apply table-wide border settings
        set_table_borders(table)
        
        # Set column headers
        headers = ['Age', 'You', 'You + spouse', 'You + 1 child', 'You + 2 children',
                   'You + 3 (or more) children', 'You + spouse + 1 child',
                   'You + spouse + 2 children', 'You + spouse + 3 (or more) children']
        
        # Format header row
        header_row = table.rows[0]
        for col, header_text in enumerate(headers):
            cell = header_row.cells[col]
            cell.text = header_text
            
            # Gray background for headers (#b7b7b7)
            set_cell_background(cell, 'B7B7B7')
            
            # Set borders for header cells
            if col == 0:  # First column (Age header) - needs left and top outer borders
                set_cell_borders(cell, top='CCCCCC', bottom='CCCCCC', left='CCCCCC', right='EDEFF6')
            elif col == 8:  # Last column - needs right and top outer borders
                set_cell_borders(cell, top='CCCCCC', bottom='CCCCCC', left='EDEFF6', right='CCCCCC')
            else:  # Middle columns - just top border for outer edge
                set_cell_borders(cell, top='CCCCCC', bottom='CCCCCC', left='EDEFF6', right='EDEFF6')
            
            # Format text
            para = cell.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para.space_before = Pt(0)
            para.space_after = Pt(0)
            
            for run in para.runs:
                run.font.bold = True
                run.font.size = Pt(9)
                run.font.name = 'Calibri'
            
            # Set vertical alignment
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        
        # Fill in ages and data
        ages = list(range(18, 64)) + ['64+']
        
        # Get allowance data for this state/class
        state_allowances = allowance_data.get(state_name, {})
        
        for row_idx, age in enumerate(ages, start=1):
            if row_idx < len(table.rows):
                # Age cell
                age_cell = table.rows[row_idx].cells[0]
                age_cell.text = str(age)
                
                # Light purple background for age column (#e0e3fe)
                set_cell_background(age_cell, 'E0E3FE')
                
                # Set borders for age cells
                if row_idx == 1:  # First data row - no top border between header
                    set_cell_borders(age_cell, bottom='EDEFF6', left='CCCCCC', right='CCCCCC')
                elif row_idx == len(ages):  # Last row - needs bottom outer border
                    set_cell_borders(age_cell, top='EDEFF6', bottom='CCCCCC', left='CCCCCC', right='CCCCCC')
                else:  # Middle rows
                    set_cell_borders(age_cell, top='EDEFF6', bottom='EDEFF6', left='CCCCCC', right='CCCCCC')
                
                # Format text
                para = age_cell.paragraphs[0]
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                para.space_before = Pt(0)
                para.space_after = Pt(0)
                
                for run in para.runs:
                    run.font.bold = True
                    run.font.size = Pt(9)
                    run.font.name = 'Calibri'
                
                # Set vertical alignment
                age_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                
                # Get allowance values for this age
                age_key = str(age)
                age_allowances = state_allowances.get(age_key, {})
                
                # Column mapping
                column_mapping = {
                    1: 'EE',   # You
                    2: 'ES',   # You + spouse
                    3: 'EC1',  # You + 1 child
                    4: 'EC2',  # You + 2 children
                    5: 'ECmax', # You + 3 (or more) children
                    6: 'FA1',  # You + spouse + 1 child
                    7: 'FA2',  # You + spouse + 2 children
                    8: 'FAmax' # You + spouse + 3 (or more) children
                }
                
                # Format data cells in this row
                for col_idx in range(1, 9):
                    data_cell = table.rows[row_idx].cells[col_idx]
                    
                    # Set allowance value if available
                    if age_allowances and col_idx in column_mapping:
                        value = age_allowances.get(column_mapping[col_idx], '')
                        data_cell.text = format_currency(value)
                    
                    # Set borders for data cells - all #cccccc
                    set_cell_borders(data_cell, top='CCCCCC', bottom='CCCCCC', left='CCCCCC', right='CCCCCC')
                    
                    para = data_cell.paragraphs[0]
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    para.space_before = Pt(0)
                    para.space_after = Pt(0)
                    
                    for run in para.runs:
                        run.font.size = Pt(9)
                        run.font.name = 'Calibri'
                    
                    # Set row height to be more compact
                    tc = data_cell._tc
                    tcPr = tc.get_or_add_tcPr()
                    tcH = OxmlElement('w:tcH')
                    tcH.set(qn('w:val'), '240')  # Smaller height in twips
                    tcH.set(qn('w:hRule'), 'atLeast')
                    tcPr.append(tcH)
        
        # Set column widths
        for row in table.rows:
            # Age column
            row.cells[0].width = Inches(0.6)
            # Data columns - wider to fill landscape page
            for i in range(1, 9):
                row.cells[i].width = Inches(1.15)
    
    # Create tables for each state - each on its own page
    for state in states:
        # Create anchor name - replace spaces and special chars with underscores
        anchor_name = state.lower().replace(' ', '_').replace('-', '_')
        create_state_table(state, anchor_name, add_page_break=True)
    
    # Add minimal final spacing
    final_spacing = doc.add_paragraph()
    final_spacing.space_before = Pt(12)
    
    # Save the document
    doc.save(output_filename)
    print(f"Document created successfully: {output_filename}")
    print("\nNotes:")
    print("- Page size: 11\" x 8.5\" (Letter, Landscape)")
    print("- Header uses the provided image directly")
    if states:
        print(f"- Created tables for {len(states)} class(es): {', '.join(states)}")
    else:
        print("- No class tables created (no classes provided)")
    print("- Table headers have gray background (#B7B7B7)")
    print("- Age columns have light purple background (#E0E3FE)")
    print("- Hyperlinks properly navigate to respective sections")
    if csv_path:
        print("- Allowance data populated from CSV")

# Usage examples
if __name__ == "__main__":
    # Example 1: Create document with CSV data
    create_ichra_document(
        output_filename='ICHRA_Allowance_Model_2025.docx',
        header_image_path='zorro_header.png',
        csv_path='allowances.csv'  # Will auto-detect classes from CSV
    )
    
    # Example 2: Create document with manual states (no CSV data)
    # create_ichra_document(
    #     output_filename='ICHRA_Model_Manual.docx',
    #     header_image_path='zorro_header.png',
    #     states=['CA', 'MA', 'CO']
    # )
    
    # Example 3: Create empty template
    # create_ichra_document(
    #     output_filename='ICHRA_Model_Template.docx',
    #     header_image_path='zorro_header.png',
    #     states=[]
    # )